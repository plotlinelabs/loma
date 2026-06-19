"""Service Agreement Review Tool.

Provides CLI commands for downloading, reading, annotating, and uploading
Word (.docx) service agreements with tracked changes and comments.

Commands:
  1. agreement_review.py download --file-id ID --user-email E --auth-token T
     Downloads a .docx from Google Drive to /tmp for processing.

  2. agreement_review.py read --file-path PATH
     Reads a .docx and returns all paragraphs as JSON.

  3. agreement_review.py annotate --file-path PATH --output-path OUT
     Reads a JSON payload from stdin with changes and comments to apply.
     Applies tracked changes and comments to the document and saves.

  4. agreement_review.py upload --file-path PATH --user-email E --auth-token T [--folder-id FID] [--name N]
     Uploads the annotated .docx to Google Drive and returns the link.

Usage (called by the agent via Bash):
  python3 tools/agreement_review.py download --file-id FILE_ID --user-email user@example.com --auth-token TOKEN
  python3 tools/agreement_review.py read --file-path /tmp/agreement.docx
  cat changes.json | python3 tools/agreement_review.py annotate --file-path /tmp/agreement.docx --output-path /tmp/agreement_reviewed.docx
  python3 tools/agreement_review.py upload --file-path /tmp/agreement_reviewed.docx --user-email user@example.com --auth-token TOKEN

Requires: python-docx>=1.2.0
"""

import asyncio
import argparse
import json
import os
import sys
import logging
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def _tracked_delete_run(run_element, author: str, rev_id: int):
    """Wrap an existing <w:r> element in a <w:del> tracked-change container."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    parent = run_element.getparent()
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), str(rev_id))
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))

    del_run = deepcopy(run_element)
    for t in del_run.findall(qn("w:t")):
        t.tag = qn("w:delText")
    del_elem.append(del_run)

    idx = list(parent).index(run_element)
    parent.remove(run_element)
    parent.insert(idx, del_elem)
    return del_elem


def _tracked_insert_text(para_element, text, after_element, author, rev_id, rpr_element=None):
    """Insert a <w:ins> tracked-change container after *after_element*."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ins_elem = OxmlElement("w:ins")
    ins_elem.set(qn("w:id"), str(rev_id))
    ins_elem.set(qn("w:author"), author)
    ins_elem.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))

    run_elem = OxmlElement("w:r")
    if rpr_element is not None:
        run_elem.insert(0, deepcopy(rpr_element))
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    run_elem.append(t_elem)
    ins_elem.append(run_elem)

    idx = list(para_element).index(after_element) + 1
    para_element.insert(idx, ins_elem)
    return ins_elem


async def cmd_download(args):
    """Download a .docx from Google Drive."""
    sys.path.insert(0, os.path.dirname(__file__))
    from _google_auth import get_google_access_token
    import aiohttp

    try:
        token = await get_google_access_token(args.user_email)
    except Exception as e:
        return {"error": f"Auth failed: {e}"}

    headers = {"Authorization": f"Bearer {token}"}

    meta_url = f"https://www.googleapis.com/drive/v3/files/{args.file_id}?fields=name,mimeType"
    async with aiohttp.ClientSession() as session:
        async with session.get(meta_url, headers=headers) as resp:
            if resp.status != 200:
                text = await resp.text()
                return {"error": f"Failed to get file metadata (HTTP {resp.status}): {text[:300]}"}
            meta = await resp.json()

    file_name = meta.get("name", "agreement.docx")
    mime = meta.get("mimeType", "")

    if "google-apps" in mime and "document" in mime:
        dl_url = f"https://www.googleapis.com/drive/v3/files/{args.file_id}/export?mimeType=application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        dl_url = f"https://www.googleapis.com/drive/v3/files/{args.file_id}?alt=media"

    async with aiohttp.ClientSession() as session:
        async with session.get(dl_url, headers=headers) as resp:
            if resp.status != 200:
                text = await resp.text()
                return {"error": f"Failed to download file (HTTP {resp.status}): {text[:300]}"}
            content = await resp.read()

    if not file_name.endswith(".docx"):
        file_name = file_name.rsplit(".", 1)[0] + ".docx"

    output_path = f"/tmp/{file_name}"
    with open(output_path, "wb") as f:
        f.write(content)

    return {
        "file_path": output_path,
        "file_name": file_name,
        "size_bytes": len(content),
    }


def cmd_read(args):
    """Read a .docx and return paragraphs as JSON."""
    from docx import Document

    try:
        doc = Document(args.file_path)
    except Exception as e:
        return {"error": f"Failed to open document: {e}"}

    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "text": text,
                "style": para.style.name if para.style else None,
            })

    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append({"row_index": ri, "cells": cells})
        tables.append({"table_index": ti, "rows": rows})

    return {
        "file_path": args.file_path,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def cmd_annotate(args):
    """Apply tracked changes and comments to a .docx based on JSON from stdin."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        payload = json.load(sys.stdin)
    except json.JSONDecodeError as e:
        return {"error": f"Invalid JSON from stdin: {e}"}

    author = payload.get("author", "Loma Legal Review")
    changes = payload.get("changes", [])
    comments = payload.get("comments", [])

    try:
        doc = Document(args.file_path)
    except Exception as e:
        return {"error": f"Failed to open document: {e}"}

    rev_id_counter = 1000
    changes_applied = 0
    comments_added = 0
    errors = []

    for ci, change in enumerate(changes):
        change_type = change.get("type", "replace")
        para_idx = change.get("paragraph_index")

        if para_idx is None or para_idx >= len(doc.paragraphs):
            errors.append(f"Change {ci}: invalid paragraph_index {para_idx}")
            continue

        para = doc.paragraphs[para_idx]

        if change_type == "replace":
            find_text = change.get("find", "")
            replace_text = change.get("replace_with", "")
            if not find_text:
                errors.append(f"Change {ci}: missing 'find' text")
                continue

            applied = False
            for run in para.runs:
                if find_text in run.text:
                    original_text = run.text
                    before = original_text.split(find_text, 1)[0]
                    after_text = original_text.split(find_text, 1)[1]
                    rpr = run._element.find(qn("w:rPr"))

                    p_elem = para._element
                    run_idx = list(p_elem).index(run._element)
                    p_elem.remove(run._element)
                    insert_pos = run_idx

                    if before:
                        before_run = OxmlElement("w:r")
                        if rpr is not None:
                            before_run.insert(0, deepcopy(rpr))
                        before_t = OxmlElement("w:t")
                        before_t.text = before
                        before_t.set(qn("xml:space"), "preserve")
                        before_run.append(before_t)
                        p_elem.insert(insert_pos, before_run)
                        insert_pos += 1

                    del_elem = OxmlElement("w:del")
                    del_elem.set(qn("w:id"), str(rev_id_counter))
                    rev_id_counter += 1
                    del_elem.set(qn("w:author"), author)
                    del_elem.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
                    del_run = OxmlElement("w:r")
                    if rpr is not None:
                        del_run.insert(0, deepcopy(rpr))
                    del_t = OxmlElement("w:delText")
                    del_t.text = find_text
                    del_t.set(qn("xml:space"), "preserve")
                    del_run.append(del_t)
                    del_elem.append(del_run)
                    p_elem.insert(insert_pos, del_elem)
                    insert_pos += 1

                    ins_elem = OxmlElement("w:ins")
                    ins_elem.set(qn("w:id"), str(rev_id_counter))
                    rev_id_counter += 1
                    ins_elem.set(qn("w:author"), author)
                    ins_elem.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
                    ins_run = OxmlElement("w:r")
                    if rpr is not None:
                        ins_run.insert(0, deepcopy(rpr))
                    ins_t = OxmlElement("w:t")
                    ins_t.text = replace_text
                    ins_t.set(qn("xml:space"), "preserve")
                    ins_run.append(ins_t)
                    ins_elem.append(ins_run)
                    p_elem.insert(insert_pos, ins_elem)
                    insert_pos += 1

                    if after_text:
                        after_run = OxmlElement("w:r")
                        if rpr is not None:
                            after_run.insert(0, deepcopy(rpr))
                        after_t = OxmlElement("w:t")
                        after_t.text = after_text
                        after_t.set(qn("xml:space"), "preserve")
                        after_run.append(after_t)
                        p_elem.insert(insert_pos, after_run)

                    applied = True
                    changes_applied += 1
                    break

            if not applied:
                full_text = para.text
                if find_text in full_text:
                    errors.append(
                        f"Change {ci}: text '{find_text[:50]}' found in paragraph but spans multiple XML runs. "
                        f"Try using a shorter substring."
                    )
                else:
                    errors.append(f"Change {ci}: text '{find_text[:50]}' not found in paragraph {para_idx}")

        elif change_type == "delete":
            find_text = change.get("find", "")
            if not find_text:
                errors.append(f"Change {ci}: missing 'find' text for delete")
                continue

            applied = False
            for run in para.runs:
                if find_text in run.text:
                    original_text = run.text
                    before = original_text.split(find_text, 1)[0]
                    after_text = original_text.split(find_text, 1)[1]
                    rpr = run._element.find(qn("w:rPr"))
                    p_elem = para._element
                    run_idx = list(p_elem).index(run._element)
                    p_elem.remove(run._element)
                    insert_pos = run_idx

                    if before:
                        before_run = OxmlElement("w:r")
                        if rpr is not None:
                            before_run.insert(0, deepcopy(rpr))
                        before_t = OxmlElement("w:t")
                        before_t.text = before
                        before_t.set(qn("xml:space"), "preserve")
                        before_run.append(before_t)
                        p_elem.insert(insert_pos, before_run)
                        insert_pos += 1

                    del_elem = OxmlElement("w:del")
                    del_elem.set(qn("w:id"), str(rev_id_counter))
                    rev_id_counter += 1
                    del_elem.set(qn("w:author"), author)
                    del_elem.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
                    del_run = OxmlElement("w:r")
                    if rpr is not None:
                        del_run.insert(0, deepcopy(rpr))
                    del_t = OxmlElement("w:delText")
                    del_t.text = find_text
                    del_t.set(qn("xml:space"), "preserve")
                    del_run.append(del_t)
                    del_elem.append(del_run)
                    p_elem.insert(insert_pos, del_elem)
                    insert_pos += 1

                    if after_text:
                        after_run = OxmlElement("w:r")
                        if rpr is not None:
                            after_run.insert(0, deepcopy(rpr))
                        after_t = OxmlElement("w:t")
                        after_t.text = after_text
                        after_t.set(qn("xml:space"), "preserve")
                        after_run.append(after_t)
                        p_elem.insert(insert_pos, after_run)

                    applied = True
                    changes_applied += 1
                    break

            if not applied:
                errors.append(f"Change {ci}: text '{find_text[:50]}' not found in paragraph {para_idx}")

        elif change_type == "insert":
            after_text = change.get("after_text", "")
            insert_text = change.get("insert_text", "")
            if not insert_text:
                errors.append(f"Change {ci}: missing 'insert_text'")
                continue

            if after_text:
                applied = False
                for run in para.runs:
                    if after_text in run.text:
                        rpr = run._element.find(qn("w:rPr"))
                        _tracked_insert_text(
                            para._element, insert_text, run._element,
                            author, rev_id_counter, rpr
                        )
                        rev_id_counter += 1
                        applied = True
                        changes_applied += 1
                        break
                if not applied:
                    errors.append(f"Change {ci}: after_text '{after_text[:50]}' not found")
            else:
                children = list(para._element)
                if children:
                    rpr_source = para.runs[-1]._element.find(qn("w:rPr")) if para.runs else None
                    _tracked_insert_text(
                        para._element, insert_text, children[-1],
                        author, rev_id_counter, rpr_source
                    )
                    rev_id_counter += 1
                    changes_applied += 1
                else:
                    errors.append(f"Change {ci}: paragraph {para_idx} is empty, cannot insert")

    for ci, comment_data in enumerate(comments):
        para_idx = comment_data.get("paragraph_index")
        comment_text = comment_data.get("comment", "")
        anchor_run_idx = comment_data.get("anchor_run_index", 0)

        if para_idx is None or para_idx >= len(doc.paragraphs):
            errors.append(f"Comment {ci}: invalid paragraph_index {para_idx}")
            continue

        para = doc.paragraphs[para_idx]
        if not para.runs:
            errors.append(f"Comment {ci}: paragraph {para_idx} has no runs to anchor to")
            continue

        if anchor_run_idx >= len(para.runs):
            anchor_run_idx = 0

        try:
            target_runs = para.runs[anchor_run_idx]
            doc.add_comment(
                runs=target_runs,
                text=comment_text,
                author=author,
                initials="GLR",
            )
            comments_added += 1
        except Exception as e:
            errors.append(f"Comment {ci}: failed to add comment: {e}")

    output_path = args.output_path or args.file_path.replace(".docx", "_reviewed.docx")
    try:
        doc.save(output_path)
    except Exception as e:
        return {"error": f"Failed to save document: {e}"}

    result = {
        "output_path": output_path,
        "changes_applied": changes_applied,
        "comments_added": comments_added,
    }
    if errors:
        result["errors"] = errors
    return result


async def cmd_upload(args):
    """Upload a .docx to Google Drive and return the link."""
    sys.path.insert(0, os.path.dirname(__file__))
    from _google_auth import get_google_access_token
    import aiohttp

    try:
        token = await get_google_access_token(args.user_email)
    except Exception as e:
        return {"error": f"Auth failed: {e}"}

    file_path = Path(args.file_path)
    if not file_path.exists():
        return {"error": f"File not found: {args.file_path}"}

    file_name = args.name or file_path.name
    content = file_path.read_bytes()

    boundary = "----FormBoundary7MA4YWxkTrZu0gW"
    metadata = {"name": file_name, "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    if args.folder_id:
        metadata["parents"] = [args.folder_id]

    body_parts = []
    body_parts.append(f"--{boundary}\r\n")
    body_parts.append("Content-Type: application/json; charset=UTF-8\r\n\r\n")
    body_parts.append(json.dumps(metadata) + "\r\n")
    body_parts.append(f"--{boundary}\r\n")
    body_parts.append("Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n")

    body_prefix = "".join(body_parts).encode("utf-8")
    body_suffix = f"\r\n--{boundary}--\r\n".encode("utf-8")
    body = body_prefix + content + body_suffix

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": f"multipart/related; boundary={boundary}",
    }

    upload_url = "https://www.googleapis.com/upload/drive/v3/files?uploadType=multipart&fields=id,name,webViewLink"

    async with aiohttp.ClientSession() as session:
        async with session.post(upload_url, headers=headers, data=body, timeout=aiohttp.ClientTimeout(total=60)) as resp:
            if resp.status not in (200, 201):
                text = await resp.text()
                return {"error": f"Upload failed (HTTP {resp.status}): {text[:300]}"}
            result = await resp.json()

    return {
        "file_id": result.get("id"),
        "file_name": result.get("name"),
        "web_view_link": result.get("webViewLink"),
    }


def main():
    from dotenv import load_dotenv
    load_dotenv()

    parser = argparse.ArgumentParser(description="Service Agreement Review Tool")
    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    dl = subparsers.add_parser("download", help="Download a .docx from Google Drive")
    dl.add_argument("--file-id", required=True, help="Google Drive file ID")
    dl.add_argument("--user-email", required=True)
    dl.add_argument("--auth-token", required=True)

    rd = subparsers.add_parser("read", help="Read a .docx and return paragraphs as JSON")
    rd.add_argument("--file-path", required=True, help="Path to the .docx file")

    an = subparsers.add_parser("annotate", help="Apply tracked changes and comments from stdin JSON")
    an.add_argument("--file-path", required=True, help="Path to the .docx file")
    an.add_argument("--output-path", required=False, help="Output path (defaults to *_reviewed.docx)")

    up = subparsers.add_parser("upload", help="Upload a .docx to Google Drive")
    up.add_argument("--file-path", required=True, help="Path to the .docx file")
    up.add_argument("--user-email", required=True)
    up.add_argument("--auth-token", required=True)
    up.add_argument("--folder-id", required=False, help="Google Drive folder ID")
    up.add_argument("--name", required=False, help="File name on Drive")

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        print(json.dumps({"error": "No command specified"}))
        sys.exit(1)

    if args.command == "download":
        result = asyncio.run(cmd_download(args))
    elif args.command == "read":
        result = cmd_read(args)
    elif args.command == "annotate":
        result = cmd_annotate(args)
    elif args.command == "upload":
        result = asyncio.run(cmd_upload(args))
    else:
        result = {"error": f"Unknown command: {args.command}"}

    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
